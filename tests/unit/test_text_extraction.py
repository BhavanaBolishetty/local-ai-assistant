"""Tests for per-file-type text extraction."""

from io import BytesIO

import pymupdf
import pytest
from docx import Document

from src.rag.text_extraction import extract_text


def _make_pdf_bytes(text: str) -> bytes:
    with pymupdf.open() as document:
        page = document.new_page()
        page.insert_text((72, 72), text)
        return document.tobytes()


def test_extracts_plain_text_for_txt_and_md() -> None:
    assert extract_text("notes.txt", b"hello world") == "hello world"
    assert extract_text("notes.md", b"# Title") == "# Title"


def test_invalid_utf8_raises_value_error() -> None:
    with pytest.raises(ValueError):
        extract_text("notes.txt", b"\xff\xfe\x00\x01")


def test_unsupported_extension_raises_value_error() -> None:
    with pytest.raises(ValueError):
        extract_text("archive.zip", b"whatever")


def test_extracts_text_from_docx() -> None:
    document = Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text("notes.docx", buffer.getvalue())

    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_pdf_extraction_extracts_real_text() -> None:
    data = _make_pdf_bytes("Hello from a real PDF page.")

    text = extract_text("notes.pdf", data)

    assert "Hello from a real PDF page." in text


def test_pdf_extraction_recovers_from_missing_trailer() -> None:
    """Reproduces the real-world "Stream has ended unexpectedly" pypdf
    failure (a valid PDF missing its trailing %%EOF/trailer) — PyMuPDF
    should recover the content directly, no special-casing needed."""
    data = _make_pdf_bytes("Recoverable despite a stripped trailer.")
    truncated = data[: data.rfind(b"%%EOF")]

    text = extract_text("notes.pdf", truncated)

    assert "Recoverable despite a stripped trailer." in text


def test_pdf_extraction_recovers_from_missing_xref() -> None:
    """Reproduces the deeper "startxref not found" pypdf failure (the
    xref table itself is gone, not just the trailer)."""
    data = _make_pdf_bytes("Recoverable despite a missing xref table.")
    truncated = data[: data.rfind(b"xref")]

    text = extract_text("notes.pdf", truncated)

    assert "Recoverable despite a missing xref table." in text


def test_pdf_extraction_raises_value_error_on_garbage() -> None:
    with pytest.raises(ValueError):
        extract_text("notes.pdf", b"this is not a pdf at all")


def test_pdf_extraction_follows_visual_order_not_stream_order() -> None:
    """Some PDFs (templated resumes, multi-column layouts) write content
    to the file in an order that doesn't match how it's laid out on the
    page. Extraction should follow visual (top-to-bottom) order, not
    whatever order the content happened to be written in."""
    with pymupdf.open() as document:
        page = document.new_page()
        # Write the visually-lower text to the content stream FIRST, and
        # the visually-higher text SECOND — deliberately out of order, the
        # way some PDF generators lay out their content stream.
        page.insert_text((72, 300), "LOWER SECTION")
        page.insert_text((72, 72), "UPPER SECTION")
        data = document.tobytes()

    text = extract_text("resume.pdf", data)

    assert text.index("UPPER SECTION") < text.index("LOWER SECTION")


def test_replacement_characters_are_normalized_to_a_hyphen() -> None:
    """Some PDF fonts encode dashes with glyph IDs that don't map to a
    real Unicode codepoint, extracting as U+FFFD (e.g. a date range like
    "Jan 2023 <?> Jul 2024"). Normalize it to a plain, readable hyphen."""
    assert extract_text("notes.txt", "Jan 2023 � Jul 2024".encode()) == "Jan 2023 - Jul 2024"
